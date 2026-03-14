import streamlit as st
import PyPDF2
from docx import Document
import json
import os
import time
import io
from dotenv import load_dotenv
from google import genai
from google.genai import types

# --- โหลดตั้งค่าเริ่มต้น ---
load_dotenv()

def get_config(key):
    try:
        if key in st.secrets: return st.secrets[key]
    except Exception: pass
    return os.getenv(key)

API_KEY = get_config("GEMINI_API_KEY")
ADMIN_PW = get_config("ADMIN_PASSWORD")

if not API_KEY:
    st.error("เกิดข้อผิดพลาด: ไม่พบ API KEY สำหรับ AI")
    client = None
else:
    client = genai.Client(api_key=API_KEY)

# --- ระบบจัดการส่วนกลาง (Global Storage) ---
DATA_FILE = "server_source_data.txt"
CONFIG_FILE = "server_config.json"

def save_admin_data(text_content, num_q, custom_inst, welcome_msg):
    with open(DATA_FILE, "w", encoding="utf-8") as f: f.write(text_content)
    config = {"num_q": num_q, "custom_inst": custom_inst, "welcome_msg": welcome_msg}
    with open(CONFIG_FILE, "w", encoding="utf-8") as f: json.dump(config, f)

def get_admin_config():
    if os.path.exists(CONFIG_FILE):
        with open(CONFIG_FILE, "r", encoding="utf-8") as f: return json.load(f)
    return None

def extract_text(uploaded_file):
    f_type = uploaded_file.name.split('.')[-1].lower()
    text = ""
    try:
        if f_type == 'pdf':
            reader = PyPDF2.PdfReader(uploaded_file)
            for page in reader.pages: text += page.extract_text()
        elif f_type == 'docx':
            doc = Document(uploaded_file)
            for para in doc.paragraphs: text += para.text + "\n"
    except Exception as e: st.error(f"อ่านไฟล์พลาด: {e}")
    return text

# --- ฟังก์ชันสร้างไฟล์ DOCX สำหรับดาวน์โหลด ---
def create_export_file(quiz_data):
    doc = Document()
    doc.add_heading('Mockup Exam เก็งแบบทดสอบ', 0)
    
    # ส่วนที่ 1: โจทย์
    doc.add_heading('ส่วนที่ 1: คำถาม', level=1)
    for i, q in enumerate(quiz_data):
        doc.add_paragraph(f"ข้อที่ {i+1}: {q['q']}")
        if q['type'] == 'choice':
            for opt in q['options']:
                doc.add_paragraph(f"    [  ] {opt}")
        elif q['type'] == 'short':
            doc.add_paragraph("    ตอบ: ____________________________________________________\n")
        else:
            doc.add_paragraph("    วิธีทำ/คำอธิบาย:")
            doc.add_paragraph("\n\n\n")
            
    # ขึ้นหน้าใหม่สำหรับเฉลย
    doc.add_page_break()
    
    # ส่วนที่ 2: เฉลย
    doc.add_heading('ส่วนที่ 2: เฉลยและคำอธิบาย', level=1)
    for i, q in enumerate(quiz_data):
        p = doc.add_paragraph()
        p.add_run(f"ข้อที่ {i+1}:\n").bold = True
        p.add_run(f"เฉลย: {q['a']}\n")
        p.add_run(f"คำอธิบาย: {q['detail']}")
        doc.add_paragraph("-" * 30)
        
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- ฟังก์ชันเจนข้อสอบ (User) ---
def user_generate_quiz():
    if not os.path.exists(DATA_FILE) or not os.path.exists(CONFIG_FILE):
        return False, "ไม่พบข้อมูลส่วนกลาง กรุณาแจ้งแอดมินให้อัปโหลดไฟล์"
        
    with open(DATA_FILE, "r", encoding="utf-8") as f: content = f.read()
    config = get_admin_config()
    
    final_instruction = f"""
    คุณคืออาจารย์ผู้เชี่ยวชาญ หน้าที่คือสร้างข้อสอบ {config['num_q']} ข้อ 
    จากเนื้อหาที่กำหนด ในรูปแบบ JSON เท่านั้น
    ข้อกำหนด: คละประเภท (choice, short, long) 
    คำสั่งพิเศษ: {config['custom_inst']}
    รูปแบบ JSON:
    [
      {{"type": "choice", "q": "โจทย์", "options": ["A","B","C","D"], "a": "คำตอบ", "detail": "คำอธิบาย"}},
      {{"type": "short", "q": "โจทย์เติมคำ", "a": "คำตอบ", "detail": "คำอธิบาย"}},
      {{"type": "long", "q": "โจทย์อัตนัย/คำนวณ", "a": "วิธีทำละเอียด", "detail": "หลักการ"}}
    ]
    """
    try:
        response = client.models.generate_content(
            model='gemini-2.5-flash',
            contents=f"เนื้อหา: {content[:30000]}",
            config=types.GenerateContentConfig(system_instruction=final_instruction, temperature=0.8)
        )
        clean_json = response.text.replace('```json', '').replace('```', '').strip()
        st.session_state.quiz_data = json.loads(clean_json)
        st.session_state.current_idx = 0
        st.session_state.user_answers = {}
        st.session_state.app_mode = "start"
        return True, "สร้างข้อสอบสำเร็จ!"
    except Exception as e:
        return False, f"AI เจนข้อสอบพลาด: {e}"

# --- ตั้งค่า Session State ---
if 'quiz_data' not in st.session_state: st.session_state.quiz_data = []
if 'current_idx' not in st.session_state: st.session_state.current_idx = 0
if 'app_mode' not in st.session_state: st.session_state.app_mode = "start"
if 'user_answers' not in st.session_state: st.session_state.user_answers = {}
if 'login_attempts' not in st.session_state: st.session_state.login_attempts = 0
if 'history' not in st.session_state: st.session_state.history = [] 

# --- UI หน้าเว็บ ---
st.set_page_config(page_title="AI Exam Self-Service", layout="centered")

# --- ส่วนแอดมิน ---
with st.sidebar:
    st.header("🔐 Admin System")
    if st.session_state.login_attempts >= 3:
        st.error("🚨 ระบบล็อก!")
    else:
        pw = st.text_input("รหัสผ่าน", type="password")
        if pw:
            if pw == ADMIN_PW:
                st.session_state.login_attempts = 0 
                st.success("เข้าสู่ระบบแอดมิน")
                st.divider()
                w_msg = st.text_area("คำอธิบายหน้าเว็บ", value="เตรียมพร้อมก่อนสอบ! กดสร้างข้อสอบเพื่อเริ่มทบทวนความรู้ได้เลย")
                n_q = st.number_input("จำนวนข้อสอบต่อชุด", min_value=1, max_value=20, value=5)
                c_inst = st.text_area("คำสั่ง AI", value="เน้นวิเคราะห์และคำนวณ")
                file = st.file_uploader("อัปโหลดเนื้อหา (PDF/Word)", type=["pdf", "docx"])
                
                if st.button("💾 บันทึกเนื้อหาและเปิดระบบ") and file:
                    with st.spinner("กำลังบันทึกข้อมูลเข้าระบบกลาง..."):
                        text = extract_text(file)
                        save_admin_data(text, n_q, c_inst, w_msg)
                        st.success("บันทึกระบบส่วนกลางเรียบร้อย!")
            else:
                st.session_state.login_attempts += 1
                st.error("รหัสผิด!")
                time.sleep(2)

# --- ส่วนการใช้งานของนักศึกษา ---
global_config = get_admin_config()

st.title("🎓 ระบบซ้อมสอบอัจฉริยะ")
if global_config:
    st.markdown(f"*{global_config['welcome_msg']}*")
st.divider()

if not global_config:
    st.info("🚧 ระบบยังไม่เปิดให้บริการ รอแอดมินอัปโหลดเนื้อหาส่วนกลางสักครู่นะครับ")
else:
    # สถานะ 1: หน้าแรก 
    if st.session_state.app_mode == "start":
        if len(st.session_state.quiz_data) == 0:
            st.warning("คุณยังไม่มีชุดข้อสอบในระบบของคุณ")
            if st.button("✨ สุ่มข้อสอบชุดใหม่ให้ฉัน (Generate Quiz)"):
                with st.spinner("AI กำลังสร้างข้อสอบเฉพาะสำหรับคุณ..."):
                    success, msg = user_generate_quiz()
                    if success: st.rerun()
                    else: st.error(msg)
        else:
            st.success(f"📌 มีข้อสอบที่ค้างอยู่ {len(st.session_state.quiz_data)} ข้อ")
            col1, col2 = st.columns(2)
            with col1:
                if st.button("🚀 ลุยเลย (เริ่มทำแบบทดสอบ)"):
                    st.session_state.app_mode = "quiz"
                    st.rerun()
            with col2:
                export_buffer = create_export_file(st.session_state.quiz_data)
                st.download_button(
                    label="📥 โหลดชีทข้อสอบ (พร้อมเฉลย)",
                    data=export_buffer,
                    file_name="Exam_Practice.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

        if len(st.session_state.history) > 0:
            st.divider()
            st.subheader("📚 ประวัติการทำข้อสอบของคุณ")
            st.caption("คลิกเพื่อทบทวนข้อสอบชุดเก่าที่คุณทำไปแล้วในรอบนี้")
            
            for i, record in enumerate(reversed(st.session_state.history)):
                quiz_num = len(st.session_state.history) - i
                with st.expander(f"📝 ข้อสอบชุดที่ {quiz_num} (ส่งเมื่อ: {record['time']})"):
                    for j, q in enumerate(record['quiz']):
                        st.markdown(f"**ข้อ {j+1}: {q['q']}**")
                        old_ans = record['ans'].get(f"q_{j}", "ไม่ได้ตอบ")
                        st.write(f"✍️ **คำตอบของคุณ:** {old_ans}")
                        st.success(f"✅ **เฉลย:** {q['a']}")
                        st.info(f"💡 **คำอธิบาย:** {q['detail']}")
                        st.markdown("---")

    # สถานะ 2: กำลังทำข้อสอบ
    elif st.session_state.app_mode == "quiz":
        idx = st.session_state.current_idx
        q = st.session_state.quiz_data[idx]
        
        st.progress((idx + 1) / len(st.session_state.quiz_data))
        st.write(f"**ข้อที่ {idx + 1} จาก {len(st.session_state.quiz_data)}**")
        st.markdown(f"### {q['q']}")

        ans_key = f"q_{idx}"
        if q['type'] == "choice":
            st.session_state.user_answers[ans_key] = st.radio("เลือกคำตอบ:", q['options'], key=ans_key)
        elif q['type'] == "short":
            st.session_state.user_answers[ans_key] = st.text_input("พิมพ์คำตอบของคุณ:", key=ans_key)
        else:
            st.session_state.user_answers[ans_key] = st.text_area("เขียนวิธีทำ/คำอธิบาย:", key=ans_key, height=150)

        st.write("") 
        col1, col2 = st.columns(2)
        with col1:
            if idx > 0:
                if st.button("⬅️ ข้อก่อนหน้า"):
                    st.session_state.current_idx -= 1
                    st.rerun()
        with col2:
            if idx < len(st.session_state.quiz_data) - 1:
                if st.button("ข้อถัดไป ➡️"):
                    st.session_state.current_idx += 1
                    st.rerun()
            else:
                if st.button("✅ ตรวจคำตอบ (ส่งข้อสอบ)"):
                    st.session_state.history.append({
                        "quiz": list(st.session_state.quiz_data),
                        "ans": dict(st.session_state.user_answers),
                        "time": time.strftime("%H:%M:%S") 
                    })
                    st.session_state.app_mode = "result"
                    st.rerun()

    # สถานะ 3: ดูเฉลย
    elif st.session_state.app_mode == "result":
        st.subheader("📊 ผลลัพธ์และเฉลยละเอียด")
        for i, q in enumerate(st.session_state.quiz_data):
            with st.expander(f"ข้อที่ {i+1}: {q['q'][:60]}..."):
                u_ans = st.session_state.user_answers.get(f"q_{i}", "ไม่ได้ตอบ")
                st.write(f"**คุณตอบ:** {u_ans}")
                st.success(f"**เฉลย:** {q['a']}")
                st.info(f"**คำอธิบาย:** {q['detail']}")
        
        st.divider()
        st.write("### ทำข้อสอบชุดนี้เสร็จแล้ว! เอาอีกไหม?")
        col1, col2 = st.columns(2)
        with col1:
            if st.button("🔄 กลับไปหน้าหลัก (ดูประวัติ)"):
                st.session_state.app_mode = "start"
                st.session_state.current_idx = 0
                st.rerun()
        with col2:
            if st.button("✨ ท้าทายตัวเอง: สุ่มข้อสอบชุดใหม่เอี่ยม!"):
                with st.spinner("AI กำลังสร้างข้อสอบชุดใหม่ให้คุณ..."):
                    success, msg = user_generate_quiz()
                    if success: st.rerun()
                    else: st.error(msg)