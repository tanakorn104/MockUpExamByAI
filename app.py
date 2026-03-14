import streamlit as st
import json
import os
from google import genai
from google.genai import types
from docx import Document
import io
import datetime

# ==========================================
# 0. การตั้งค่าและโหลด API Key
# ==========================================
def get_api_key():
    try:
        return st.secrets["GEMINI_API_KEY"]
    except:
        return os.getenv("GEMINI_API_KEY")

API_KEY = get_api_key()

def load_web_config():
    try:
        with open("web_config.json", "r", encoding="utf-8") as f:
            return json.load(f)
    except FileNotFoundError:
        return {
            "page_title": "ระบบเก็งข้อสอบ Computer Architecture",
            "welcome_message": "ยินดีต้อนรับสู่ระบบจำลองข้อสอบ (บทที่ 7-8)",
            "button_text": "🚀 เริ่มสุ่มสร้างข้อสอบชุดใหม่",
            "success_message": "สร้างข้อสอบสำเร็จ! ขอให้โชคดีในการทำ"
        }

config = load_web_config()
st.set_page_config(page_title=config.get("page_title"), page_icon="🎓")

if not API_KEY:
    st.error("❌ ไม่พบ API Key กรุณาตรวจสอบการตั้งค่า")
    st.stop()

client = genai.Client(api_key=API_KEY)

# ==========================================
# 1. ระบบจัดการประวัติ (History State)
# ==========================================
if "exam_history" not in st.session_state:
    st.session_state.exam_history = []

# ==========================================
# 2. ฟังก์ชันสร้างไฟล์ Word แบบหลายโหมด
# ==========================================
def create_docx(quiz_data=None, user_answers=None, mode="worksheet", history_data=None):
    doc = Document()
    
    # --- โหมดใบงานก่อนทำ (โจทย์อยู่หน้าแรก เฉลยอยู่หน้าหลังสุด) ---
    if mode == "worksheet":
        doc.add_heading('ใบงานข้อสอบจำลอง (Worksheet)', 0)
        # ส่วนที่ 1: โจทย์ล้วนๆ
        for i, q in enumerate(quiz_data):
            doc.add_heading(f"ข้อที่ {i+1}: {q.get('q', 'ไม่พบโจทย์')}", level=1)
            if str(q.get('type')).upper() == 'CHOICE':
                for opt in q.get('options', []):
                    doc.add_paragraph(f"  [ ] {opt}")
            else:
                doc.add_paragraph("  ..................................................................")
            doc.add_paragraph("")
            
        # ขึ้นหน้าใหม่สำหรับเฉลย
        doc.add_page_break()
        doc.add_heading('เฉลยข้อสอบ (Answer Key)', 0)
        for i, q in enumerate(quiz_data):
            doc.add_heading(f"ข้อที่ {i+1}: {q.get('q', '')}", level=2)
            doc.add_paragraph(f"เฉลย: {q.get('a', '')}", style='Intense Quote')
            doc.add_paragraph(f"คำอธิบาย: {q.get('detail', '')}")
            doc.add_paragraph("-" * 20)

    # --- โหมดสรุปผล (หลังจากทำเสร็จ 1 ชุด) ---
    elif mode == "result":
        doc.add_heading('สรุปผลข้อสอบ', 0)
        score = 0
        for i, q in enumerate(quiz_data):
            u_ans = str(user_answers.get(i, "")).strip().lower()
            correct_ans = str(q.get('a', '')).strip().lower()
            if u_ans == correct_ans and u_ans != "":
                score += 1
                
        doc.add_paragraph(f"คะแนนที่ทำได้: {score} / {len(quiz_data)} คะแนน")
        doc.add_paragraph("=" * 30)

        for i, q in enumerate(quiz_data):
            doc.add_heading(f"ข้อที่ {i+1}: {q.get('q', '')}", level=1)
            doc.add_paragraph(f"คำตอบของคุณ: {user_answers.get(i, 'ไม่ได้ตอบ')}")
            doc.add_paragraph(f"เฉลย: {q.get('a', '')}", style='Intense Quote')
            doc.add_paragraph(f"คำอธิบาย: {q.get('detail', '')}")
            doc.add_paragraph("-" * 20)

    # --- โหมดโหลดประวัติทั้งหมด (รวมทุกชุดที่เคยทำ) ---
    elif mode == "history":
        doc.add_heading('ประวัติการทำข้อสอบทั้งหมด (Exam History)', 0)
        doc.add_paragraph(f"วันที่พิมพ์เอกสาร: {datetime.datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
        
        for idx, record in enumerate(history_data):
            doc.add_heading(f"ชุดข้อสอบที่ {idx+1} | คะแนน: {record['score']} / {record['total']}", level=1)
            for i, q in enumerate(record['quiz_data']):
                doc.add_paragraph(f"ข้อ {i+1}: {q.get('q', '')}", style='List Number')
                doc.add_paragraph(f"คุณตอบ: {record['user_answers'].get(i, 'ไม่ได้ตอบ')}")
                doc.add_paragraph(f"เฉลย: {q.get('a', '')}", style='Intense Quote')
                doc.add_paragraph(f"คำอธิบาย: {q.get('detail', '')}\n")
            doc.add_page_break()

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

# ==========================================
# 3. ฟังก์ชันเรียก AI เจนข้อสอบ
# ==========================================
def generate_quiz():
    try:
        with open("pre_processed_content.txt", "r", encoding="utf-8") as f:
            content = f.read()
        with open("instruction.txt", "r", encoding="utf-8") as f:
            final_instruction = f.read()

        json_format_prompt = """
        สำคัญมาก: กรุณาตอบกลับเป็นรูปแบบ JSON Array เท่านั้น ห้ามมีข้อความอธิบายอื่นปะปน 
        
        !!! กฎการสร้างข้อสอบ (ต้องทำตามอย่างเคร่งครัดเพื่อให้ได้ระดับเดียวกับข้อสอบมหาวิทยาลัย) !!!
        1. ให้ออกข้อสอบจำนวนทั้งหมด 24 ข้อ
        2. โจทย์คำถาม (q) ต้องมีความยาว เป็น "สถานการณ์จำลอง" (Scenario-based) มีการกำหนดค่าตัวแปร อ้างอิงสถาปัตยกรรม หรือสั่งให้วิเคราะห์/คำนวณ ห้ามถามสั้นๆ ทื่อๆ เด็ดขาด
        3. ข้อ 1-10 ให้กำหนด "type": "CHOICE": โดย "options" ทั้ง 4 ตัวเลือกต้องเป็นประโยคอธิบายยาวๆ เชิงลึก
        4. ข้อ 11-20 ให้กำหนด "type": "SHORT": เป็นข้อเขียนเชิงวิเคราะห์ อธิบายหลักการ หรือเปรียบเทียบ
        5. ข้อ 21-24 ให้กำหนด "type": "SHORT": เป็นโจทย์สถานการณ์ขนาดใหญ่ โดยในโจทย์ (q) ต้องแบ่งเป็น "Part A:" และ "Part B:" อย่างชัดเจน (เช่น การคำนวณ CPU Utilization, ออกแบบ Instruction Format)
        6. ฟิลด์ "detail" ต้องอธิบายเฉลยยาวๆ ทีละขั้นตอน (Step-by-step) อย่างละเอียดที่สุด

        ตัวอย่างโครงสร้าง JSON:
        [
            {
                "type": "CHOICE",
                "q": "โจทย์สถานการณ์จำลองที่ยาวและซับซ้อน...",
                "options": ["ตัวเลือกยาวและละเอียด 1", "ตัวเลือกยาวและละเอียด 2", "ตัวเลือกยาว 3", "ตัวเลือก 4"],
                "a": "คำตอบที่ถูกต้องตรงกับใน options",
                "detail": "คำอธิบายละเอียดเชิงลึก ทำไมถูก ทำไมผิด..."
            },
            {
                "type": "SHORT",
                "q": "สถานการณ์ยาวมาก... Part A: ให้วิเคราะห์... Part B: ให้ออกแบบ...",
                "options": [],
                "a": "เฉลยแบบเจาะลึก...",
                "detail": "คำอธิบายขั้นตอนอย่างละเอียด..."
            }
        ]
        """
        
        full_prompt = f"{final_instruction}\n{json_format_prompt}\n\nเนื้อหา:\n{content}"

        response = client.models.generate_content(
            model='gemini-3.1-flash-lite',
            contents=full_prompt,
            config=types.GenerateContentConfig(
                temperature=0.8,
                response_mime_type="application/json" # นำกลับมาใช้ได้เพราะเป็นโมเดล Gemini
            )
        )
        
        # ทำความสะอาดข้อความเผื่อ AI ใส่ tag markdown มาด้วย
        res_text = response.text.strip()
        if res_text.startswith("```json"):
            res_text = res_text[7:]
        elif res_text.startswith("```"):
            res_text = res_text[3:]
        if res_text.endswith("```"):
            res_text = res_text[:-3]
        res_text = res_text.strip()
        
        st.session_state.quiz_data = json.loads(res_text, strict=False)
        st.session_state.user_answers = {}
        st.session_state.app_mode = "quiz_running"
        return True
    except Exception as e:
        st.error(f"เกิดข้อผิดพลาด: {e}")
        return False

# ==========================================
# 4. แถบเมนูด้านข้าง (Sidebar) สำหรับประวัติ
# ==========================================
with st.sidebar:
    st.header("📂 เก็บประวัติลงเครื่อง")
    st.info("ระบบจะจำข้อสอบที่คุณทำในหน้าเว็บนี้ หากต้องการเก็บไว้อ่านถาวร ให้กดดาวน์โหลดก่อนปิดเว็บนะครับ")
    
    if len(st.session_state.exam_history) > 0:
        st.success(f"คุณทำข้อสอบไปแล้ว {len(st.session_state.exam_history)} ชุด")
        
        # ปุ่มโหลดประวัติเป็นไฟล์ Word
        history_docx = create_docx(mode="history", history_data=st.session_state.exam_history)
        st.download_button(
            label="💾 โหลดประวัติทั้งหมด (Word)",
            data=history_docx,
            file_name="ComArch_All_History.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )
        
        # ปุ่มโหลดประวัติเป็นไฟล์ JSON เผื่อใช้ทำ Data
        history_json = json.dumps(st.session_state.exam_history, ensure_ascii=False, indent=4)
        st.download_button(
            label="📄 โหลดประวัติรูปแบบข้อมูล (JSON)",
            data=history_json,
            file_name="ComArch_History.json",
            mime="application/json",
            use_container_width=True
        )
    else:
        st.write("ยังไม่มีประวัติการทำข้อสอบ")

# ==========================================
# 5. หน้าจอหลัก (UI)
# ==========================================
if "app_mode" not in st.session_state:
    st.session_state.app_mode = "start"

st.title(f"🎓 {config.get('page_title')}")

# --- หน้าแรก ---
if st.session_state.app_mode == "start":
    st.markdown(f"**{config.get('welcome_message')}**")
    if st.button(config.get('button_text'), use_container_width=True):
        with st.spinner("AI กำลังสร้างข้อสอบ..."):
            if generate_quiz():
                st.rerun()

# --- หน้าทำข้อสอบ ---
elif st.session_state.app_mode == "quiz_running":
    st.success(config.get('success_message'))
    
    # ดาวน์โหลดใบงาน (มีโจทย์ด้านหน้า เฉลยอยู่หน้าหลัง)
    worksheet_data = create_docx(quiz_data=st.session_state.quiz_data, mode="worksheet")
    st.download_button(
        label="📥 ดาวน์โหลดใบงาน (โจทย์อยู่ด้านหน้า เฉลยอยู่ด้านหลัง)",
        data=worksheet_data,
        file_name="ComArch_Worksheet_with_Key.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True
    )
    st.divider()

    with st.form("exam_form"):
        temp_answers = {}
        for i, q in enumerate(st.session_state.quiz_data):
            st.markdown(f"**ข้อที่ {i+1}:** {q.get('q')}")
            q_type = str(q.get('type')).upper()
            
            if q_type == "CHOICE":
                temp_answers[i] = st.radio(f"คำตอบข้อ {i+1}", q.get('options', []), key=f"ans_{i}", index=None, label_visibility="collapsed")
            else:
                temp_answers[i] = st.text_input(f"คำตอบข้อ {i+1}", key=f"ans_{i}", placeholder="พิมพ์คำตอบสั้นๆ...", label_visibility="collapsed")
            st.write("") 

        if st.form_submit_button("📤 ส่งข้อสอบและตรวจคำตอบ", use_container_width=True):
            st.session_state.user_answers = temp_answers
            
            # คำนวณคะแนนเพื่อเก็บลงประวัติ
            score = sum(1 for idx, q in enumerate(st.session_state.quiz_data) 
                        if str(temp_answers.get(idx, "")).strip().lower() == str(q.get('a', '')).strip().lower())
            
            # บันทึกประวัติลง Session State
            st.session_state.exam_history.append({
                "quiz_data": st.session_state.quiz_data,
                "user_answers": temp_answers,
                "score": score,
                "total": len(st.session_state.quiz_data)
            })
            
            st.session_state.app_mode = "result"
            st.rerun()

# --- หน้าผลลัพธ์ ---
elif st.session_state.app_mode == "result":
    # คำนวณคะแนนปัจจุบัน
    score = sum(1 for i, q in enumerate(st.session_state.quiz_data) 
                if str(st.session_state.user_answers.get(i, "")).strip().lower() == str(q.get('a', '')).strip().lower())
    total_q = len(st.session_state.quiz_data)
    
    st.header(f"🎯 คะแนนของคุณ: {score} / {total_q}")
    st.progress(score / total_q if total_q > 0 else 0)

    for i, q in enumerate(st.session_state.quiz_data):
        u_ans = st.session_state.user_answers.get(i)
        is_correct = str(u_ans).strip().lower() == str(q.get('a')).strip().lower()
        
        with st.container(border=True):
            st.write(f"**ข้อที่ {i+1}:** {q.get('q')}")
            if is_correct:
                st.success(f"✅ ถูกต้อง! (คำตอบของคุณ: {u_ans})")
            else:
                st.error(f"❌ ผิด (คำตอบของคุณ: {u_ans})")
                st.write(f"**เฉลยที่ถูกต้อง:** {q.get('a')}")
            st.info(f"💡 **คำอธิบาย:** {q.get('detail')}")

    # ปุ่มดาวน์โหลดเฉลยชุดปัจจุบัน และ เริ่มใหม่
    col1, col2 = st.columns(2)
    with col1:
        result_data = create_docx(quiz_data=st.session_state.quiz_data, user_answers=st.session_state.user_answers, mode="result")
        st.download_button("💾 โหลดผลลัพธ์ชุดนี้", result_data, "Exam_Result_Current.docx", use_container_width=True)
    with col2:
        if st.button("🔄 สุ่มสร้างชุดใหม่", use_container_width=True):
            st.session_state.app_mode = "start"
            st.rerun()